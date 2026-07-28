"""Section-by-section content of the lab report.

Kept apart from `build_report.py` so that file stays a thin set of docx helpers
and this one reads as prose. Every number comes from the helpers' live reads of
`artifacts/`, so a retrain refreshes the report.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

import config
from build_report import (
    BEST,
    CATEGORICAL,
    METRICS,
    NUMERIC,
    RANKING,
    SCHEMA,
    add_bullets,
    add_comment,
    add_equation,
    add_figure,
    add_heading,
    add_numbered,
    add_para,
    add_screenshot_slot,
    add_table,
    top_features,
)

SECTIONS = [
    "Introduction",
    "Objectives",
    "Problem Statement",
    "Dataset Description",
    "Software and Hardware Requirements",
    "Machine Learning Models Used",
    "Methodology",
    "Data Preprocessing",
    "Model Training",
    "Model Evaluation",
    "User Input",
    "Prediction",
    "Feature Importance and Interpretability",
    "User Interface",
    "Visualization",
    "Sample Output",
    "Deployment",
    "Advantages",
    "Limitations",
    "Future Improvements",
    "Result",
    "Conclusion",
]

money = lambda v: f"${v:,.0f}"


def cover_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    for text, size, bold in [
        ("SYLHET ENGINEERING COLLEGE", 20, True),
        ("Department of Computer Science & Engineering", 14, False),
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold, run.font.size = bold, Pt(size)

    doc.add_paragraph()
    add_comment(doc, "Replace every [BRACKETED] value with your own details, then delete "
                     "all red NOTE TO AUTHOR lines before submitting.")

    add_table(doc, ["Field", "Value"], [
        ["Course Name", "Introduction To Data Science Sessional"],
        ["Course No", "CSE 604"],
        ["Project Title", "House Price Prediction using Multiple Regression Models"],
        ["Date of Submission", "[DD.MM.YYYY]"],
        ["Remarks", ""],
    ], widths=[2.0, 4.3])

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run("Submitted To").bold = True
    for line in ["[Instructor Name]", "[Designation]", "Sylhet Engineering College"]:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(line)

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run("Submitted By").bold = True

    add_table(doc, ["Name", "Registration No.", "Session"], [
        ["[Your Name]", "[Registration No.]", "[Session]"],
        ["[Group Member]", "[Registration No.]", "[Session]"],
    ], widths=[2.6, 2.0, 1.7])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def table_of_contents(doc):
    add_heading(doc, "Table of Contents", level=1)
    add_comment(doc, "In Word, select this list and use References → Table of Contents "
                     "to generate one with live page numbers.")
    for index, title in enumerate(SECTIONS, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.add_run(f"{index}. {title}")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def introduction(doc):
    add_heading(doc, "1. Introduction", level=1)
    add_para(doc,
        "Machine Learning (ML) is a branch of Artificial Intelligence that enables computers to "
        "learn patterns from data and make predictions without being explicitly programmed for "
        "every task. Where a classification model predicts a discrete label, a regression model "
        "predicts a continuous numeric quantity — which is the problem addressed in this project.")
    add_para(doc,
        f"The Ames Housing dataset is used to develop a machine learning regression system. "
        f"{len(RANKING)} different regression algorithms are trained on identical data: Linear "
        "Regression, Random Forest, Gradient Boosting, Support Vector Regression, a Voting "
        "Ensemble and a Stacking Ensemble.")
    add_para(doc,
        "The models are compared using four standard regression metrics — MAE, MSE, RMSE and R² "
        "— measured on the same hold-out test split. A web application built with a FastAPI "
        "backend and a React frontend lets a user enter the characteristics of a house and see "
        "the price predicted by every model alongside that model's measured accuracy.")
    add_para(doc,
        "The project additionally includes automated exploratory data analysis, permutation-based "
        "feature importance for interpretability, and diagnostic plots showing where the models "
        "succeed and where they fail.")


def objectives(doc):
    add_heading(doc, "2. Objectives", level=1)
    add_bullets(doc, [
        "To understand the complete workflow of a machine learning regression problem.",
        "To perform exploratory data analysis on a real-world dataset before modelling.",
        "To build a preprocessing pipeline handling missing values, feature scaling and "
        "categorical encoding.",
        f"To train {len(RANKING)} regression models on the same dataset and the same split.",
        "To compare the models using MAE, MSE, RMSE and R².",
        "To combine individual models into Voting and Stacking ensembles.",
        "To predict the sale price of a house from user-provided characteristics.",
        "To interpret the models by measuring the contribution of each input feature.",
        "To evaluate model behaviour using residual and actual-vs-predicted diagnostics.",
        "To develop an interactive web interface and deploy it online.",
    ])


def problem_statement(doc):
    add_heading(doc, "3. Problem Statement", level=1)
    add_para(doc,
        "Develop a machine learning application that trains multiple regression models on the "
        "Ames Housing dataset and predicts the sale price of a residential property.")
    add_para(doc, "The application should allow a user to provide the following characteristics:")
    add_bullets(doc, [config.FEATURE_LABELS.get(f["name"], f["name"]) for f in SCHEMA])
    add_para(doc, "The system should display:")
    add_bullets(doc, [
        "The predicted sale price from each model",
        "The measured accuracy (R²) of each model",
        "The best-performing model, clearly highlighted",
        "The average prediction across all models",
        "A visual comparison of MAE, MSE, RMSE and R²",
        "The importance of each input feature",
        "Diagnostic plots of prediction error",
        "An automated exploratory data analysis report",
    ])
    add_para(doc, "The interface should be a modern web application deployable to a public URL.")


def dataset_description(doc):
    add_heading(doc, "4. Dataset Description", level=1)
    add_para(doc,
        "The Ames Housing dataset describes residential property sales in Ames, Iowa, USA. It is "
        "a standard benchmark for regression and forms the basis of the Kaggle competition "
        "“House Prices: Advanced Regression Techniques”. The published dataset contains "
        "1,460 records described by 81 columns. For this project 14 interpretable columns were "
        "selected so that a user can realistically fill them into a form, while still covering "
        "the main drivers of price: size, quality, age and location.")

    add_heading(doc, "4.1 Dataset Summary", level=2)
    add_table(doc, ["Property", "Value"], [
        ["Source", "Kaggle / OpenML, fetched via scikit-learn"],
        ["Total records", "1,460"],
        ["Selected features", f"{len(SCHEMA)} ({len(NUMERIC)} numeric, {len(CATEGORICAL)} categorical)"],
        ["Target variable", "SalePrice (continuous, US dollars)"],
        ["Target range", "$34,900 – $755,000 (median $163,000)"],
        ["Missing values", "81 (all in GarageType)"],
        ["Train / test split", "80% / 20% — 1,168 / 292 records"],
        ["Encoded model inputs", "57 columns after preprocessing"],
    ], widths=[2.3, 4.0])

    add_heading(doc, "4.2 Numeric Features", level=2)
    add_table(doc, ["Column", "Description", "Range"],
        [[f["name"], config.FEATURE_LABELS.get(f["name"], f["name"]),
          f"{f['minimum']:g} – {f['maximum']:g}"] for f in NUMERIC],
        widths=[1.5, 3.0, 1.8])

    add_heading(doc, "4.3 Categorical Features", level=2)
    add_table(doc, ["Column", "Description", "Levels"],
        [[f["name"], config.FEATURE_LABELS.get(f["name"], f["name"]), str(len(f["choices"]))]
         for f in CATEGORICAL],
        widths=[1.5, 3.4, 1.4])

    add_comment(doc, "Optional: add a screenshot of the raw CSV opened in Excel here to show "
                     "the data in its original form.")


def requirements(doc):
    add_heading(doc, "5. Software and Hardware Requirements", level=1)
    add_heading(doc, "5.1 Software Requirements", level=2)
    add_bullets(doc, [
        "Operating System: Windows 11",
        "Python 3.11 and Node.js 22",
        "Visual Studio Code",
        "Backend: FastAPI, Uvicorn, Pydantic",
        "Machine learning: scikit-learn 1.7.1, pandas, NumPy, joblib",
        "Frontend: React 19, Vite, Recharts",
        "Exploratory data analysis: ydata-profiling",
        "Version control and deployment: Git, GitHub, Render, Vercel",
    ])
    add_heading(doc, "5.2 Hardware Requirements", level=2)
    add_bullets(doc, [
        "Processor: Intel Core i3 or higher",
        "RAM: minimum 4 GB (8 GB recommended — training the ensembles is memory-intensive)",
        "Storage: minimum 500 MB free space",
        "Internet connection for installing packages and for deployment",
    ])
