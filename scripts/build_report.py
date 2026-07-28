"""Generate the editable Word lab report for the project.

    python scripts/build_report_figures.py    # diagrams first
    python scripts/build_report.py

Writes `report/ML_Lab_Report.docx`. Every metric in the document is read from
`artifacts/` at build time, so re-running after a retrain refreshes the numbers
instead of leaving stale ones in the text.

Placeholders the student fills in are written as [SQUARE BRACKETS]; screenshot
slots are yellow-highlighted paragraphs.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

import config

ROOT = Path(__file__).resolve().parent.parent
FIGURES = ROOT / "report" / "figures"
OUTPUT = ROOT / "report" / "ML_Lab_Report.docx"

ACCENT = RGBColor(0x1A, 0x7F, 0x37)
MUTED = RGBColor(0x52, 0x51, 0x4E)
PLACEHOLDER = RGBColor(0xC6, 0x28, 0x28)


# --- Live numbers -------------------------------------------------------------
METRICS: dict = json.loads(config.METRICS_FILE.read_text(encoding="utf-8"))
IMPORTANCES: dict = json.loads(config.IMPORTANCES_FILE.read_text(encoding="utf-8"))
SCHEMA: list = json.loads(config.SCHEMA_FILE.read_text(encoding="utf-8"))["features"]

RANKING = sorted(METRICS, key=lambda n: METRICS[n]["R2"], reverse=True)
BEST = RANKING[0]
NUMERIC = [f for f in SCHEMA if f["kind"] == "number"]
CATEGORICAL = [f for f in SCHEMA if f["kind"] == "category"]


def top_features(model: str, count: int) -> list[tuple[str, float]]:
    rows = IMPORTANCES[model]
    total = sum(max(r["importance"], 0.0) for r in rows) or 1.0
    return [
        (config.FEATURE_LABELS.get(r["feature"], r["feature"]),
         max(r["importance"], 0.0) / total * 100)
        for r in rows[:count]
    ]


# --- Document helpers ---------------------------------------------------------
def add_heading(doc, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT if level == 1 else MUTED
    return heading


def add_para(doc, text: str = "", bold: bool = False, italic: bool = False, size: int = 11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.size = Pt(size)
    return para


def add_bullets(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_equation(doc, text: str, note: str = ""):
    """Equations are plain text in a monospace face so they stay editable.

    python-docx cannot emit OMML, and an image would not be editable — which is
    the whole point of handing over a .docx.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.bold = True
    if note:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = caption.add_run(note)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = MUTED


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(10)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()
    return table


def add_screenshot_slot(doc, number: int, title: str, instruction: str):
    """A visually obvious placeholder the student replaces with an image."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[ SCREENSHOT {number}: {title} ]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PLACEHOLDER

    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrun = hint.add_run(f"How to capture: {instruction}")
    hrun.italic = True
    hrun.font.size = Pt(9)
    hrun.font.color.rgb = MUTED

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = caption.add_run(f"Figure {number}: {title}")
    crun.italic = True
    crun.font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc, filename: str, number: int, caption: str, width: float = 6.3):
    path = FIGURES / filename
    if not path.exists():
        add_para(doc, f"[ missing figure: {filename} — run scripts/build_report_figures.py ]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Figure {number}: {caption}")
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_comment(doc, text: str):
    """An instruction to the student; delete before submitting."""
    para = doc.add_paragraph()
    run = para.add_run(f"➤ NOTE TO AUTHOR: {text}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = PLACEHOLDER
